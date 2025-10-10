import streamlit as st
from docx import Document
import re
import math

# -----------------------
# Helpers: chuẩn hoá chuỗi
# -----------------------
def clean_text(s: str) -> str:
    return re.sub(r'\s+', ' ', s).strip()

# -----------------------
# Hàm đọc lawbank (có số thứ tự và Ref)
# -----------------------
def load_lawbank(docx_file):
    """
    Đọc lawbank: file có đánh số 1., 2., ... mỗi block có câu hỏi, đáp án a./b./c./d. 
    Đáp án đúng có dấu * ở trước chữ cái (ví dụ *a. ...). Có thể có "Ref."
    """
    try:
        doc = Document(docx_file)
    except Exception as e:
        st.error(f"❌ Không thể đọc file lawbank: {e}")
        return []

    text = "\n".join([p.text for p in doc.paragraphs if p.text and p.text.strip()])
    if not text.strip():
        return []

    # Tách blocks theo số thứ tự (1., 2., ...)
    blocks = re.finditer(r'\d+\.\s*(.*?)(?=(?:\n\s*\d+\.\s*)|\Z)', text, flags=re.S)
    questions = []

    for b in blocks:
        part = b.group(1).strip()
        if not part:
            continue

        # Loại bỏ phần Ref nếu có (không cần)
        part = re.split(r'\bRef[:.]', part, flags=re.I)[0].strip()

        # Thêm newline trước mọi marker đáp án (a. b. c. d. hoặc a) b) ...), kể cả *a.)
        part = re.sub(r'(?<!\n)(?=[*]?\s*[A-Da-d]\s*[.\)])', '\n', part)

        lines = [l.strip() for l in part.splitlines() if l.strip()]
        if not lines:
            continue

        # Dòng đầu là question (có thể dài)
        qtext = clean_text(lines[0])

        opts = []
        correct = ""

        for ln in lines[1:]:
            m = re.match(r'^[*]?\s*([A-Da-d])\s*[.\)]\s*(.*)$', ln, flags=re.S)
            if m:
                letter = m.group(1).lower()
                opt_body = clean_text(m.group(2))
                opt_string = f"{letter}. {opt_body}"
                opts.append(opt_string)
                if ln.strip().startswith("*"):
                    correct = opt_string
            else:
                # nếu có dòng không khớp marker đáp án, có thể dòng nối tiếp của đáp án trước -> nối vào cuối option trước
                if opts:
                    opts[-1] = opts[-1] + " " + clean_text(ln)
                    opts[-1] = clean_text(opts[-1])
                else:
                    # nếu chưa có options, nối vào câu hỏi
                    qtext = qtext + " " + clean_text(ln)

        if opts:
            if not correct:
                correct = opts[0]
            questions.append({"question": qtext, "options": opts, "answer": correct})

    return questions


# -----------------------
# Hàm đọc cabbank (không đánh số, đáp án có thể dính)
# -----------------------
def load_cabbank(docx_file):
    """
    Đọc cabbank: không đánh số. Mỗi câu là một hoặc nhiều đoạn văn, sau đó có a./b./c. (có thể dính).
    Đáp án đúng có dấu * ở trước chữ cái.
    """
    try:
        doc = Document(docx_file)
    except Exception as e:
        st.error(f"❌ Không thể đọc file cabbank: {e}")
        return []

    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    if not paragraphs:
        return []

    text = "\n".join(paragraphs)

    # Thêm newline trước mọi marker đáp án (a. b. c. d. hoặc a) b) ...), kể cả *a.)
    text = re.sub(r'(?<!\n)(?=[*]?\s*[A-Da-d]\s*[.\)])', '\n', text)

    lines = [l.strip() for l in text.splitlines() if l.strip()]

    questions = []
    current_q = {"question": "", "options": [], "answer": ""}

    for ln in lines:
        m = re.match(r'^[*]?\s*([A-Da-d])\s*[.\)]\s*(.*)$', ln, flags=re.S)
        if m:
            # Đây là dòng đáp án
            letter = m.group(1).lower()
            body = clean_text(m.group(2))
            opt_string = f"{letter}. {body}"
            current_q["options"].append(opt_string)
            if ln.strip().startswith("*"):
                current_q["answer"] = opt_string
        else:
            # Không phải đáp án => câu hỏi (hoặc nối tiếp câu hỏi)
            if current_q["options"]:
                # đã có options từ câu trước -> lưu câu trước và bắt đầu câu mới
                if not current_q["answer"] and current_q["options"]:
                    current_q["answer"] = current_q["options"][0]
                questions.append(current_q)
                current_q = {"question": clean_text(ln), "options": [], "answer": ""}
            else:
                # chưa có options: nối tiếp câu hỏi
                if current_q["question"]:
                    current_q["question"] = clean_text(current_q["question"] + " " + ln)
                else:
                    current_q["question"] = clean_text(ln)

    # Thêm câu cuối cùng nếu có
    if current_q["question"] and current_q["options"]:
        if not current_q["answer"]:
            current_q["answer"] = current_q["options"][0]
        questions.append(current_q)

    return questions


# -----------------------
# Giao diện Streamlit
# -----------------------
st.set_page_config(page_title="Ngân hàng câu hỏi", layout="wide")
st.title("📚 Ngân hàng câu hỏi (Lawbank & Cabbank)")

# Cho phép upload .docx để test (ưu tiên file upload nếu có)
uploaded_file = st.file_uploader("Upload file .docx (nếu muốn test file bất kỳ)", type=["docx"])

bank_choice = st.selectbox("Chọn ngân hàng muốn làm:", ["Ngân hàng Luật", "Ngân hàng Kỹ thuật"])

# Nếu upload file, dùng file đó; nếu không, dùng file mặc định
if uploaded_file is not None:
    file_source = uploaded_file
else:
    file_source = "lawbank.docx" if "Luật" in bank_choice else "cabbank.docx"

# Load questions theo lựa chọn
if "Luật" in bank_choice:
    questions = load_lawbank(file_source)
else:
    questions = load_cabbank(file_source)

# Debug: hiện số câu và 3 câu đầu trong expander để bạn kiểm tra nhanh
with st.expander("🔍 Debug: Xem trước kết quả parsing (3 câu đầu)"):
    st.write(f"Bank: {bank_choice}")
    st.write(f"Số câu đọc được: {len(questions)}")
    for idx, q in enumerate(questions[:3], start=1):
        st.markdown(f"**{idx}.** {q['question']}")
        for o in q['options']:
            marker = "✅" if o == q['answer'] else ""
            st.write(f"- {o} {marker}")

if not questions:
    st.error(f"Không đọc được câu hỏi nào từ file cho ngân hàng '{bank_choice}'.\n- Nếu bạn dùng file mặc định, kiểm tra tên file và định dạng.\n- Nếu bạn upload, đảm bảo file .docx chứa câu hỏi theo cấu trúc (câu rồi các đáp án a./b./c.).")
    st.stop()

# Phân nhóm
TOTAL = len(questions)
group_size = 10
num_groups = math.ceil(TOTAL / group_size)
group_labels = [f"Câu {i*group_size+1} - {min((i+1)*group_size, TOTAL)}" for i in range(num_groups)]

# Session state cơ bản
if "current_bank" not in st.session_state:
    st.session_state.current_bank = bank_choice
if "last_group" not in st.session_state:
    st.session_state.last_group = None
if "submitted" not in st.session_state:
    st.session_state.submitted = False

# Reset khi đổi ngân hàng
if st.session_state.current_bank != bank_choice:
    for k in list(st.session_state.keys()):
        if k.startswith("q_"):
            del st.session_state[k]
    st.session_state.submitted = False
    st.session_state.current_bank = bank_choice

# Chọn nhóm
selected_group = st.selectbox("📘 Bạn muốn làm nhóm câu nào?", group_labels, index=0)

if st.session_state.last_group != (selected_group + bank_choice):
    for k in list(st.session_state.keys()):
        if k.startswith("q_"):
            del st.session_state[k]
    st.session_state.submitted = False
    st.session_state.last_group = selected_group + bank_choice

start = group_labels.index(selected_group) * group_size
end = min(start + group_size, TOTAL)
batch = questions[start:end]

# Hiển thị câu hỏi
placeholder_choice = "-- Chưa chọn --"

if not st.session_state.submitted:
    st.markdown(f"### 🧩 Nhóm {selected_group}  (tổng {len(batch)} câu)")
    for i, q in enumerate(batch, start=start + 1):
        st.markdown(f"<div style='text-align:left'><b>{i}. {q['question']}</b></div>", unsafe_allow_html=True)
        options_ui = [placeholder_choice] + q["options"]
        # Lưu giá trị chọn (mặc định là placeholder)
        st.radio("", options_ui, index=0, key=f"q_{i}")
        st.markdown("<hr>", unsafe_allow_html=True)

    if st.button("✅ Nộp bài và xem kết quả"):
        unanswered = [i for i in range(start + 1, end + 1) if st.session_state.get(f"q_{i}") in (None, placeholder_choice)]
        if unanswered:
            st.warning(f"⚠️ Bạn chưa chọn đáp án cho {len(unanswered)} câu: {', '.join(map(str, unanswered))}")
        else:
            st.session_state.submitted = True
            st.experimental_rerun()

else:
    # Tính điểm và hiển thị kết quả
    score = 0
    for i, q in enumerate(batch, start=start + 1):
        selected = st.session_state.get(f"q_{i}")
        # Nếu user chọn placeholder, treat as None
        if selected in (None, placeholder_choice):
            selected_display = None
        else:
            selected_display = selected
        correct = q["answer"]

        if selected_display == correct:
            score += 1
            st.success(f"{i}. {q['question']}\n\n✅ Đúng — {correct}")
        else:
            st.error(f"{i}. {q['question']}\n\n❌ Sai. Bạn chọn: **{selected_display or '---'}**  → Đáp án đúng: **{correct}**")
        st.markdown("<hr>", unsafe_allow_html=True)

    st.subheader(f"🎯 Kết quả: {score}/{len(batch)} câu đúng")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔁 Làm lại nhóm này"):
            for i in range(start + 1, end + 1):
                key = f"q_{i}"
                if key in st.session_state:
                    del st.session_state[key]
            st.session_state.submitted = False
            st.experimental_rerun()
    with col2:
        if st.button("➡️ Sang nhóm khác"):
            for i in range(start + 1, end + 1):
                key = f"q_{i}"
                if key in st.session_state:
                    del st.session_state[key]
            st.session_state.submitted = False
            st.experimental_rerun()
